import os
import json
import random
import sys
from faker import Faker
from datetime import datetime, timedelta
from pymongo import MongoClient
import requests
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont

# Initialize Faker
fake = Faker()

# -------------------------------------------------
# Load Email Templates from data1.json (subject and body variations)
# -------------------------------------------------
def load_email_templates(file_path="data1.json"):
    """
    Load email templates from the provided JSON file.
    Expects a top-level key "variation_templates" mapping to various template types.
    """
    with open(file_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    return data.get("variation_templates", {})

# Global: load our templates once using data1.json.
email_templates = load_email_templates()

def generate_email_from_template(template_key, set_no, acc_no, templates):
    """
    Generate email subject and body from the given template.
    
    Parameters:
      template_key (str): Template key (e.g. "new_request", "follow_up", "update", "duplicate", etc.)
      set_no (int or str): Value to replace {set_no}.
      acc_no (int or str): Value to replace {acc_no}.
      templates (dict): Dictionary of templates loaded from JSON.
    
    Returns:
      tuple: subject and body strings with placeholders replaced.
    """
    variations = templates.get(template_key)
    if not variations:
        raise ValueError(f"Template key '{template_key}' not found in templates.")
    chosen_variation = random.choice(variations)
    subject = chosen_variation.get("subject", "").format(set_no=set_no, acc_no=acc_no)
    body = chosen_variation.get("body", "").format(set_no=set_no, acc_no=acc_no)
    return subject, body

# -------------------------------------------------
# Helper: Create realistic attachment file with sample data.
# (Unchanged from your original code)
# -------------------------------------------------
def create_dummy_attachment(filename):
    os.makedirs("attachments", exist_ok=True)
    file_path = os.path.join("attachments", filename)
    ext = filename.split(".")[-1].lower()
    
    sample_data = (
        f"Name: {fake.name()}\n"
        f"Address: {fake.address()}\n"
        f"Phone: {fake.phone_number()}\n"
        f"SSN: {fake.ssn()}\n"
        f"TIN: {fake.random_number(digits=9, fix_len=True)}\n"
        f"Primary Contact: {fake.phone_number()}\n"
        f"NAC Code: {fake.lexify(text='????')}\n"
        f"Loan Amount: {random.randint(1000, 10000)}\n"
        f"Requested Income: {random.randint(30000, 100000)}\n"
        f"Expenses: {random.randint(1000, 5000)}\n"
        f"Ownership Information: {fake.company()}\n"
    )
    
    if ext == "msg":
        msg_content = (
            f"From: {fake.email()}\n"
            f"To: {fake.email()}\n"
            f"Subject: {fake.sentence(nb_words=6)}\n"
            f"Date: {datetime.now().strftime('%a, %d %b %Y %H:%M:%S %z')}\n\n"
            f"{sample_data}"
        )
        with open(file_path, "w", encoding="utf-8") as f_out:
            f_out.write(msg_content)
    
    elif ext in ["txt", "csv", "html"]:
        with open(file_path, "w", encoding="utf-8") as f_out:
            f_out.write("Realistic Form Data:\n" + sample_data)
    
    elif ext == "pdf":
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
        except ImportError:
            sys.exit("Install reportlab via: pip install reportlab")
        form_type = random.choice(["Beneficial Ownership Form", "Entity Registration Form", "W-9 Form", "Tax Document"])
        c = canvas.Canvas(file_path, pagesize=letter)
        width, height = letter
        c.setFont("Helvetica-Bold", 16)
        c.drawCentredString(width / 2, height - 50, form_type)
        c.setFont("Helvetica", 12)
        y = height - 100
        for line in sample_data.split("\n"):
            if line.strip():
                c.drawString(50, y, line)
                y -= 20
        c.drawString(50, y-20, "Additional terms: This document is a simulation.")
        c.showPage()
        c.save()
    
    elif ext == "docx":
        try:
            from docx import Document
        except ImportError:
            sys.exit("Install python-docx via: pip install python-docx")
        form_type = random.choice(["Beneficial Ownership Form", "Entity Registration Form", "W-9 Form", "Tax Document"])
        from docx import Document
        doc = Document()
        doc.add_heading(form_type, 0)
        for line in sample_data.split("\n"):
            if line.strip():
                doc.add_paragraph(line)
        doc.save(file_path)
    
    elif ext == "xlsx":
        try:
            from openpyxl import Workbook
        except ImportError:
            sys.exit("Install openpyxl via: pip install openpyxl")
        wb = Workbook()
        ws = wb.active
        ws.title = "User Data"
        headers = ["Name", "Address", "Phone", "SSN", "TIN", "Primary Contact", "NAC Code", "Loan Amount", "Requested Income", "Expenses", "Ownership Information"]
        ws.append(headers)
        for _ in range(10):
            row = [
                fake.name(),
                fake.address(),
                fake.phone_number(),
                fake.ssn(),
                str(fake.random_number(digits=9, fix_len=True)),
                fake.phone_number(),
                fake.lexify(text="????"),
                random.randint(1000, 10000),
                random.randint(30000, 100000),
                random.randint(1000, 5000),
                fake.company()
            ]
            ws.append(row)
        wb.save(file_path)
    
    elif ext == "zip":
        import zipfile
        inner_filename = "data.csv"
        headers = ["Name", "Email", "Phone", "Address", "SSN", "TIN"]
        data_rows = [headers]
        for _ in range(5):
            data_rows.append([fake.name(), fake.email(), fake.phone_number(), fake.address(), fake.ssn(), str(fake.random_number(digits=9, fix_len=True))])
        csv_str = "\n".join([",".join(row) for row in data_rows])
        import zipfile
        with zipfile.ZipFile(file_path, 'w') as zf:
            zf.writestr(inner_filename, csv_str)
    
    elif ext in ["png", "jpg", "jpeg"]:
        width, height = 600, 400
        id_card = Image.new("RGB", (width, height), color="white")
        draw = ImageDraw.Draw(id_card)
        photo_box = (50, 50, 250, 250)
        draw.rectangle(photo_box, outline="black", width=2)
        try:
            response = requests.get("https://placekitten.com/200/200", timeout=5)
            cat_img = Image.open(BytesIO(response.content))
            cat_img = cat_img.resize((photo_box[2] - photo_box[0], photo_box[3] - photo_box[1]))
            id_card.paste(cat_img, (photo_box[0], photo_box[1]))
        except Exception:
            draw.rectangle(photo_box, fill="gray")
        text_x = 300
        text_y = 50
        try:
            font = ImageFont.truetype("arial.ttf", 20)
        except Exception:
            font = ImageFont.load_default()
        id_fields = {
            "Name": fake.name(),
            "DOB": fake.date_of_birth(minimum_age=18, maximum_age=90).strftime("%Y-%m-%d"),
            "Address": fake.address().replace("\n", ", "),
            "ID Number": str(fake.random_number(digits=8, fix_len=True)),
            "Phone": fake.phone_number()
        }
        for label, value in id_fields.items():
            draw.text((text_x, text_y), f"{label}: {value}", fill="black", font=font)
            text_y += 30
        id_card.save(file_path)
    
    else:
        with open(file_path, "w", encoding="utf-8") as f_out:
            f_out.write("Realistic Form Data:\n" + sample_data)
    
    return file_path

# -------------------------------------------------
# Generate Email Data Dump for 1000 Emails with Realistic Attachments and Updates
# -------------------------------------------------

# Define request categories and sub-requests.
request_categories = {
    "Billing Issue": ["Invoice Discrepancy", "Payment Delay", "Refund Request"],
    "Technical Support": ["Bug Report", "Feature Assistance", "Hardware Setup"],
    "Account Update": ["Profile Change", "Password Reset", "Account Deactivation"],
    "General Inquiry": ["Product Info", "Service Feedback", "Partnership Request"]
}

# Attachment file extensions (including "msg").
attachment_extensions = ["pdf", "docx", "xlsx", "csv", "zip", "txt", "png", "jpg", "html", "msg"]

def generate_attachment_filename():
    num = random.randint(1000, 9999)
    ext = random.choice(attachment_extensions)
    return f"Attachment_{num}.{ext}"

# Create recurring senders (simulate 10 recurring senders).
recurring_senders = {}
for _ in range(10):
    name = fake.name()
    email = fake.email()
    req_cat = random.choice(list(request_categories.keys()))
    sub_req = random.choice(request_categories[req_cat])
    recurring_senders[email] = {
        "sender": name,
        "request_category": req_cat,
        "sub_request": sub_req
    }

emails = []
generated_emails_indices = []

def random_datetime():
    now = datetime.now()
    delta = timedelta(days=random.randint(0, 30), hours=random.randint(0, 23), minutes=random.randint(0, 59))
    return (now - delta).isoformat()

N = 200

for i in range(N):
    # 5% chance: duplicate a previous email.
    if generated_emails_indices and random.random() < 0.10:
        dup_index = random.choice(generated_emails_indices)
        duplicate_email = emails[dup_index].copy()
        duplicate_email["email_id"] = f"EML{i+1:04d}"
        duplicate_email["received_at"] = random_datetime()
        duplicate_email["is_duplicate"] = True
        emails.append(duplicate_email)
        generated_emails_indices.append(len(emails)-1)
        continue

    email_record = {}
    email_record["email_id"] = f"EML{i+1:04d}"
    
    if random.random() < 0.15:
        sender_email = random.choice(list(recurring_senders.keys()))
        recurring_info = recurring_senders[sender_email]
        email_record["sender"] = recurring_info["sender"]
        email_record["sender_email"] = sender_email
        request_category = recurring_info["request_category"]
        sub_request = recurring_info["sub_request"]
    else:
        email_record["sender"] = fake.name()
        email_record["sender_email"] = fake.email()
        request_category = random.choice(list(request_categories.keys()))
        sub_request = random.choice(request_categories[request_category])
    
    # Instead of a simple concatenated subject/body, use our templates from data1.json.
    # We first generate random values for {set_no} and {acc_no}.
    set_no = random.randint(1, 20)   # arbitrary set number
    acc_no = fake.random_number(digits=6, fix_len=True)  # simulated account number

    # Determine which template to use based on conditions.
    if email_record.get("is_duplicate", False):
        chosen_template = "duplicate"
    elif random.random() < 0.30:  # 20% chance = update case
        chosen_template = "update"
        email_record["is_update_case"] = True
    else:
        # Optionally, randomly choose among new_request, follow_up, and general_update.
        chosen_template = random.choice(["new_request", "follow_up", "general_update"])
        email_record["is_update_case"] = False

    subject, body = generate_email_from_template(chosen_template, set_no, acc_no, email_templates)
    email_record["subject"] = subject
    email_record["body"] = body

    # For update cases, attach an extra file containing the email's data.
    is_update_case = email_record.get("is_update_case", False)
    num_attachments = random.randint(1, 3)
    attachments = [generate_attachment_filename() for _ in range(num_attachments)]
    
    if is_update_case:
        update_filename = f"Update_{email_record['email_id']}.txt"
        update_content = (
            f"Subject: {email_record['subject']}\n"
            f"Body: {email_record['body']}\n"
            f"Request: {request_category}\n"
            f"Sub-request: {sub_request}\n"
            "Note: This is an update request. Please verify and update accordingly."
        )
        os.makedirs("attachments", exist_ok=True)
        update_file_path = os.path.join("attachments", update_filename)
        with open(update_file_path, "w", encoding="utf-8") as f_out:
            f_out.write(update_content)
        attachments.append(update_filename)
    
    # Create realistic attachment files.
    for att in attachments:
        create_dummy_attachment(att)
    
    email_record["attachments"] = attachments
    email_record["received_at"] = random_datetime()
    email_record["status"] = "unprocessed"
    if "is_duplicate" not in email_record:
        email_record["is_duplicate"] = False

    emails.append(email_record)
    generated_emails_indices.append(len(emails)-1)

# -------------------------------------------------
# Insert the generated emails directly into MongoDB.
# -------------------------------------------------
try:
    client = MongoClient("mongodb://localhost:27017/")
    db = client["email_training_db"]
    collection = db["emails_training"]  # Using "emails" collection (table)
    insert_result = collection.insert_many(emails)
    print(f"Inserted {len(insert_result.inserted_ids)} emails into the 'emails' collection in MongoDB.")
except Exception as e:
    print("MongoDB insertion failed:", e)

# -------------------------------------------------
# Write out a JSON backup of the data dump.
# -------------------------------------------------
with open("sample_emails_data_dump.json", "w", encoding="utf-8") as f_out:
    json.dump(emails, f_out, indent=2)
print("Data dump JSON file 'sample_emails_data_dump.json' created.")